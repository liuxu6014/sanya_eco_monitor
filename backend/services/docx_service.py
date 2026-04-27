"""Word document report generation service.

Images are inserted INLINE at the exact sentence where （见图N） appears.
All figures are numbered sequentially starting from 图1.
Figure captions are centered beneath each image.

Fixed figure numbering is provided by the shared figure manifest built in
report_figures.py and aligned with ai_report.py.

Appendix (at end):
  附录一 虫情测报灯采集实景图像  (downloaded from InsectRecord.image_url)
  附录二 孢子捕捉仪采集实景图像  (downloaded from SporeRecord.image_url)
"""

import io
import base64
import os
import re
import logging
import subprocess
import tempfile
import time
import zipfile
from collections import defaultdict

import httpx
from PIL import Image, ImageOps
from docx import Document
from docx.shared import Inches, Pt
from docx.shared import RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from services.report_figures import (
    build_figure_manifest,
    build_figure_map as build_shared_figure_map,
)

logger = logging.getLogger(__name__)

_COVER_TITLE = "\u6a61\u80f6\u6797\u8fd1\u81ea\u7136\u5316\u6539\u9020\u751f\u6001\u6548\u76ca\u8bc4\u4f30\u62a5\u544a"
_COVER_SUBTITLE = "\u4e09\u4e9a\u5e02\u5929\u6daf\u533a\u6a61\u80f6\u6797\u8fd1\u81ea\u7136\u5316\u6539\u9020\u548c\u519c\u7530\u63d0\u5347\u76d1\u6d4b\u5e73\u53f0"
_PERIOD_PREFIX = "\u7edf\u8ba1\u5468\u671f\uff1a"
_PAGE_PREFIX = "\u7b2c "
_PAGE_SUFFIX = " \u9875"
_REMOTE_IMAGE_MAX_EDGE = 1600
_REMOTE_IMAGE_JPEG_QUALITY = 68
_REMOTE_IMAGE_SUBSAMPLING = 2
_INLINE_IMAGE_JPEG_QUALITY = 88

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _set_cjk_font(run, name='Microsoft YaHei', bold=None):
    if bold is not None:
        run.bold = bold
    run.font.name = name
    if run._element.rPr is None:
        run._element.get_or_add_rPr()
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)


def _set_run_size(run, size_pt: int | float):
    run.font.size = Pt(size_pt)


def _tighten_paragraph(paragraph, *, first_indent: bool = False, left_indent: bool = False):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing = 1.0
    if first_indent:
        fmt.first_line_indent = Pt(24)
    else:
        fmt.first_line_indent = Pt(0)
    if left_indent:
        fmt.left_indent = Pt(24)
    else:
        fmt.left_indent = Pt(0)
    return paragraph


def _set_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    _tighten_paragraph(p)
    for run in p.runs:
        _set_cjk_font(run, bold=True)
    return p


def _add_paragraph(doc, text, first_indent=True, left_indent=False):
    """Add a paragraph with CJK font and optional indent."""
    if not text or not text.strip():
        return None
    p = doc.add_paragraph()
    _tighten_paragraph(p, first_indent=first_indent, left_indent=left_indent)

    for part in re.split(r'(\*\*.*?\*\*)', text):
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            _set_cjk_font(run, bold=True)
        else:
            run = p.add_run(part)
            _set_cjk_font(run, bold=False)
    return p


def _enable_update_fields_on_open(doc: Document):
    settings = doc.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def _append_field(paragraph, instruction: str, default_text: str):
    begin_run = paragraph.add_run()
    _set_cjk_font(begin_run)
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    fld_char_begin.set(qn("w:dirty"), "true")
    begin_run._r.append(fld_char_begin)

    instr_run = paragraph.add_run()
    _set_cjk_font(instr_run)
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = f" {instruction} "
    instr_run._r.append(instr_text)

    sep_run = paragraph.add_run()
    _set_cjk_font(sep_run)
    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")
    sep_run._r.append(fld_char_sep)

    text_run = paragraph.add_run(default_text)
    _set_cjk_font(text_run)

    end_run = paragraph.add_run()
    _set_cjk_font(end_run)
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    end_run._r.append(fld_char_end)


def _remove_empty_body_paragraphs(doc: Document):
    body = doc._element.body
    for paragraph in list(doc.paragraphs):
        element = paragraph._element
        if element.getparent() != body:
            continue
        if paragraph.text.strip():
            continue
        if element.xpath(".//w:drawing"):
            continue
        if element.xpath(".//w:br"):
            continue
        if element.xpath(".//w:fldChar"):
            continue
        body.remove(element)


def _refresh_word_fields(filepath: str) -> None:
    """Use local Microsoft Word to update PAGE/NUMPAGES fields before delivery."""
    if os.name != "nt":
        return

    resolved = os.path.abspath(filepath).replace("'", "''")
    script = f"""
$path = '{resolved}'
$word = $null
$doc = $null
try {{
    $word = New-Object -ComObject Word.Application
    $word.Visible = $false
    $word.DisplayAlerts = 0
    $doc = $word.Documents.Open($path, $false, $false)
    $doc.Repaginate() | Out-Null
    $pages = $doc.ComputeStatistics(2)
    foreach ($story in $doc.StoryRanges) {{
        $range = $story
        while ($null -ne $range) {{
            $range.Fields.Update() | Out-Null
            $range = $range.NextStoryRange
        }}
    }}
    $doc.Fields.Update() | Out-Null
    $doc.Repaginate() | Out-Null
    $pages = $doc.ComputeStatistics(2)
    foreach ($story in $doc.StoryRanges) {{
        $range = $story
        while ($null -ne $range) {{
            $range.Fields.Update() | Out-Null
            $range = $range.NextStoryRange
        }}
    }}
    $doc.Save()
    Write-Output \"PAGE_COUNT=$pages\"
}} finally {{
    if ($doc -ne $null) {{ $doc.Close([ref]0) }}
    if ($word -ne $null) {{ $word.Quit() }}
}}
"""
    try:
        result = subprocess.run(
            ["powershell", "-NoProfile", "-ExecutionPolicy", "Bypass", "-Command", script],
            check=True,
            capture_output=True,
            text=True,
            timeout=180,
        )
        page_count = None
        for line in result.stdout.splitlines():
            line = line.strip()
            if not line.startswith("PAGE_COUNT="):
                continue
            try:
                page_count = int(line.split("=", 1)[1])
            except ValueError:
                page_count = None
            break
        if page_count:
            _persist_numpages_result(filepath, page_count)
    except Exception as exc:
        logger.warning("Failed to refresh Word fields for %s: %s", filepath, exc)


def _persist_numpages_result(filepath: str, page_count: int) -> None:
    page_text = str(page_count)
    with zipfile.ZipFile(filepath, "r") as src:
        footer_names = [name for name in src.namelist() if re.fullmatch(r"word/footer\d+\.xml", name)]
        if not footer_names:
            return
        entries = [(info, src.read(info.filename)) for info in src.infolist()]

    with tempfile.NamedTemporaryFile(
        delete=False,
        suffix=".docx",
        dir=os.path.dirname(os.path.abspath(filepath)),
    ) as tmp:
        temp_path = tmp.name

    try:
        with zipfile.ZipFile(temp_path, "w") as dst:
            for info, data in entries:
                if info.filename in footer_names:
                    data = _update_footer_numpages_xml(data, page_text)
                dst.writestr(info, data)
        last_error = None
        for _ in range(10):
            try:
                os.replace(temp_path, filepath)
                break
            except PermissionError as exc:
                last_error = exc
                time.sleep(0.5)
        else:
            raise last_error or PermissionError(filepath)
    finally:
        if os.path.exists(temp_path):
            try:
                os.remove(temp_path)
            except OSError:
                pass


def _update_footer_numpages_xml(xml_bytes: bytes, page_text: str) -> bytes:
    try:
        xml_text = xml_bytes.decode("utf-8")
    except UnicodeDecodeError:
        return xml_bytes

    complex_pattern = re.compile(
        r'(<w:instrText[^>]*>\s*NUMPAGES\s*</w:instrText>.*?'
        r'<w:fldChar\b[^>]*w:fldCharType="separate"[^>]*/>.*?'
        r'<w:t[^>]*>)(.*?)(</w:t>)',
        re.DOTALL,
    )
    simple_pattern = re.compile(
        r'(<w:fldSimple\b[^>]*w:instr="[^"]*NUMPAGES[^"]*"[^>]*>.*?<w:t[^>]*>)(.*?)(</w:t>)',
        re.DOTALL,
    )

    xml_text, complex_count = complex_pattern.subn(
        lambda match: f"{match.group(1)}{page_text}{match.group(3)}",
        xml_text,
        count=1,
    )
    xml_text, simple_count = simple_pattern.subn(
        lambda match: f"{match.group(1)}{page_text}{match.group(3)}",
        xml_text,
    )
    if complex_count == 0 and simple_count == 0:
        return xml_bytes
    return xml_text.encode("utf-8")


def _add_cover_title(doc, text: str):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _tighten_paragraph(paragraph)
    run = paragraph.add_run(text)
    _set_cjk_font(run, bold=True)
    _set_run_size(run, 24)
    run.font.color.rgb = RGBColor(31, 78, 121)
    return paragraph


def _normalize_image_stream(
    img_bytes: bytes,
    *,
    max_edge: int | None = None,
    jpeg_quality: int = _INLINE_IMAGE_JPEG_QUALITY,
) -> io.BytesIO:
    try:
        image = Image.open(io.BytesIO(img_bytes))
        image = ImageOps.exif_transpose(image)
        if max_edge and max(image.size) > max_edge:
            image.thumbnail((max_edge, max_edge), Image.Resampling.LANCZOS)

        output = io.BytesIO()
        if image.mode == "RGBA":
            image.save(output, format="PNG")
        else:
            if image.mode != "RGB":
                image = image.convert("RGB")
            image.save(
                output,
                format="JPEG",
                quality=jpeg_quality,
                optimize=True,
                progressive=True,
                subsampling=_REMOTE_IMAGE_SUBSAMPLING,
            )
        output.seek(0)
        return output
    except Exception:
        raw = io.BytesIO(img_bytes)
        raw.seek(0)
        return raw


def _insert_image_bytes(
    doc,
    img_bytes: bytes,
    caption: str,
    width_inches: float = 5.5,
    *,
    max_edge: int | None = None,
    jpeg_quality: int = _INLINE_IMAGE_JPEG_QUALITY,
):
    """Insert already-loaded image bytes followed by a centered caption paragraph."""
    try:
        pic_para = doc.add_paragraph()
        pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _tighten_paragraph(pic_para)
        run = pic_para.add_run()
        run.add_picture(
            _normalize_image_stream(
                img_bytes,
                max_edge=max_edge,
                jpeg_quality=jpeg_quality,
            ),
            width=Inches(width_inches),
        )
    except Exception as exc:
        logger.warning("Failed to insert image %s: %s", caption, exc)
        return

    cap_para = doc.add_paragraph()
    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _tighten_paragraph(cap_para)
    cap_run = cap_para.add_run(caption)
    _set_cjk_font(cap_run, bold=False)
    cap_run.font.size = Pt(10)


def _insert_figure(doc, image_src: str, caption: str, width_inches: float = 5.5):
    """Insert an image from either a data URI or a remote URL."""
    normalize_kwargs: dict[str, int | None] = {}
    try:
        if image_src.startswith("data:"):
            raw = image_src.split(",", 1)[-1]
            img_bytes = base64.b64decode(raw)
        elif image_src.startswith(("http://", "https://")):
            img_bytes = _download_image(image_src)
            if not img_bytes:
                return
            normalize_kwargs = {
                "max_edge": _REMOTE_IMAGE_MAX_EDGE,
                "jpeg_quality": _REMOTE_IMAGE_JPEG_QUALITY,
            }
        else:
            return
    except Exception as exc:
        logger.warning("Failed to load figure %s: %s", caption, exc)
        return

    _insert_image_bytes(doc, img_bytes, caption, width_inches=width_inches, **normalize_kwargs)


def _download_image(url: str, timeout: int = 15) -> bytes | None:
    """Download image bytes from a URL; return None on failure."""
    try:
        with httpx.Client(verify=False, timeout=timeout, follow_redirects=True) as client:
            resp = client.get(url)
            resp.raise_for_status()
            return resp.content
    except Exception as exc:
        logger.warning("Failed to download image %s: %s", url, exc)
        return None


def _insert_url_image(doc, url: str, caption: str, width_inches: float = 5.5):
    """Download image from URL and insert into the document. Handles relative URLs."""
    # Handle relative URLs if they happen to be stored as such
    if url.startswith('/'):
        # For local downloads, we might need a base URL.
        # However, many IoT URLs are absolute. If we find relative ones, we assume they are local.
        # For now, let's assume absolute or try to handle standard local paths.
        pass

    img_bytes = _download_image(url)
    if not img_bytes:
        # Insert placeholder text instead
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _tighten_paragraph(p)
        run = p.add_run(f"[图片加载失败: {caption}]")
        _set_cjk_font(run, bold=False)
        run.font.size = Pt(9)
        return

    _insert_image_bytes(
        doc,
        img_bytes,
        caption,
        width_inches=width_inches,
        max_edge=_REMOTE_IMAGE_MAX_EDGE,
        jpeg_quality=_REMOTE_IMAGE_JPEG_QUALITY,
    )


def _insert_missing_image_placeholder(doc, message: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _tighten_paragraph(p)
    run = p.add_run(f"[占位] {message}")
    _set_cjk_font(run, bold=False)
    run.font.size = Pt(10)


def _build_figure_map(
    summary: dict,
    charts: dict,
    ai_images: dict,
    figure_manifest: list[dict] | None = None,
) -> dict[int, tuple[str, str]]:
    """Build {fig_num: (image_data, caption)} from the shared figure manifest."""
    manifest = figure_manifest or build_figure_manifest(summary, charts, ai_images)
    return build_shared_figure_map(manifest)


def _value_text(value, suffix: str = "") -> str:
    if value in (None, ""):
        return "—"
    return f"{value}{suffix}"


def _append_guideline_summary(doc: Document, summary: dict) -> None:
    guideline = summary.get("guideline_metrics") or {}
    if not guideline:
        return

    methodology = guideline.get("methodology") or {}
    runoff = guideline.get("runoff_erosion") or {}
    water = guideline.get("water_quality") or {}
    pest = guideline.get("pest_management") or {}
    weather = summary.get("weather_support") or guideline.get("weather_support") or {}
    water_source_support = guideline.get("water_source_support") or {}
    implementation_matrix = guideline.get("implementation_matrix") or {}

    _set_heading(doc, "《指南》指标适配摘要", level=1)

    if methodology:
        _set_heading(doc, "一、监测体系科学性说明", level=2)
        _add_paragraph(doc, methodology.get("monitoring_statement", ""))
        _add_paragraph(doc, methodology.get("baseline_statement", ""))

    if runoff.get("available"):
        reference_station = runoff.get("reference_station") or {}
        _set_heading(doc, "二、水土流失强度减少率监测型估算", level=2)
        _add_paragraph(
            doc,
            "；".join(
                [
                    f"统计周期：最近{_value_text(runoff.get('period_days'))}天",
                    f"参照监测点：{reference_station.get('name', '—')}（{reference_station.get('device_code', '—')}）",
                    f"参照侵蚀代理指标：{_value_text(reference_station.get('erosion_proxy'))}",
                    f"其他监测点平均侵蚀代理指标：{_value_text(runoff.get('plantation_avg_proxy'))}",
                    f"估算减蚀率：{_value_text(runoff.get('estimated_reduction_rate'), '%')}",
                ]
            ),
            first_indent=False,
        )
        note = runoff.get("note")
        if note:
            _add_paragraph(doc, f"说明：{note}", first_indent=False)

    if water.get("available"):
        baseline = water.get("baseline_period") or {}
        recent = water.get("recent_period") or {}
        _set_heading(doc, "三、农业面源污染削减率", level=2)
        _add_paragraph(
            doc,
            "；".join(
                [
                    f"基准期：{_value_text(baseline.get('start'))} 至 {_value_text(baseline.get('end'))}",
                    f"基准期记录数：{_value_text(baseline.get('records_count'))}",
                    f"近30天统计期：{_value_text(recent.get('start'))} 至 {_value_text(recent.get('end'))}",
                    f"综合削减率：{_value_text(water.get('composite_reduction_rate'), '%')}",
                ]
            ),
            first_indent=False,
        )
        for item in water.get("metrics", []):
            _add_paragraph(
                doc,
                (
                    f"■ {item.get('label', '—')}：基准期平均 {_value_text(item.get('baseline_avg'))} {item.get('unit', '')}，"
                    f"近30天平均 {_value_text(item.get('recent_avg'))} {item.get('unit', '')}，"
                    f"最新值 {_value_text(item.get('latest_value'))} {item.get('unit', '')}，"
                    f"近30天削减率 {_value_text(item.get('recent_reduction_rate'), '%')}，"
                    f"最新削减率 {_value_text(item.get('latest_reduction_rate'), '%')}"
                ),
                first_indent=False,
                left_indent=True,
            )

    if pest.get("available"):
        insect_peak = pest.get("insect_peak") or {}
        spore_peak = pest.get("spore_peak") or {}
        top_species = pest.get("top_species") or {}
        _set_heading(doc, "四、虫害发生情况与适应性管理闭环", level=2)
        _add_paragraph(
            doc,
            "；".join(
                [
                    f"风险等级：{_value_text(pest.get('risk_level'))}",
                    f"主要关注虫种：{_value_text(top_species.get('name'))}",
                    f"虫情峰值：{_value_text(insect_peak.get('date'))} / {_value_text(insect_peak.get('count'))}只",
                    f"孢子峰值：{_value_text(spore_peak.get('date'))} / {_value_text(spore_peak.get('count'))}个",
                ]
            ),
            first_indent=False,
        )
        _add_paragraph(doc, pest.get("chain_text", ""), first_indent=False)
        _add_paragraph(doc, pest.get("management_record_template", ""), first_indent=False)

    if weather.get("enabled") and weather.get("status") == "ok":
        current = weather.get("current") or {}
        summary_info = weather.get("history_summary") or {}
        history_range = weather.get("history_range") or {}
        _set_heading(doc, "五、气象补充与水源涵养支撑", level=2)
        _add_paragraph(
            doc,
            "；".join(
                [
                    f"数据来源：{_value_text(weather.get('source'))}",
                    f"当前天气：{_value_text(current.get('text'))}",
                    f"当前温度：{_value_text(current.get('temp'), '℃')}",
                    f"当前湿度：{_value_text(current.get('humidity'), '%')}",
                    f"当前风速：{_value_text(current.get('wind_speed'), ' km/h')}",
                    f"历史区间：{_value_text(history_range.get('start'))} 至 {_value_text(history_range.get('end'))}",
                    f"最近7天累计降水：{_value_text(summary_info.get('total_precip'), ' mm')}",
                    f"最近7天平均气温：{_value_text(summary_info.get('avg_temp_mean'), ' ℃')}",
                    f"最近7天平均湿度：{_value_text(summary_info.get('avg_humidity'), ' %')}",
                    f"最近7天平均风速：{_value_text(summary_info.get('avg_wind_speed'), ' km/h')}",
                ]
            ),
            first_indent=False,
        )

    if water_source_support and water_source_support.get("message"):
        _add_paragraph(
            doc,
            f"水源涵养支撑说明：{water_source_support.get('message', '')}",
            first_indent=False,
        )

    confirmed_rules = implementation_matrix.get("confirmed_rules") or []
    if confirmed_rules:
        _set_heading(doc, "六、本期已确认业务口径", level=2)
        for item in confirmed_rules:
            _add_paragraph(doc, f"■ {item}", first_indent=False, left_indent=True)

    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Appendix: device capture images
# ---------------------------------------------------------------------------

def _append_device_image_section(doc, summary: dict):
    """
    Append an appendix of actual on-device capture photos from insect traps
    and spore collectors. Groups by device code, sorted by time.

    Structure:
      附录：设备采集实景图像
        一、虫情测报灯采集图像（设备编号: XXXXXX）
           [图片] 采集时间: YYYY-MM-DD HH:MM
           ...
        二、孢子捕捉仪采集图像（设备编号: XXXXXX）
           ...
    """
    insect_images: list[dict] = (summary.get("insect") or {}).get("capture_images") or []
    spore_images: list[dict] = (summary.get("spore") or {}).get("capture_images") or []

    doc.add_page_break()
    _set_heading(doc, "附录：设备采集实景图像", level=1)
    p = doc.add_paragraph()
    note_run = p.add_run("以下图像均来自各监测设备在报告周期内自动采集的实时照片，按设备分组、时间顺序排列。")
    _set_cjk_font(note_run, bold=False)
    note_run.font.size = Pt(10)

    # ---- 虫情测报灯 ----
    _set_heading(doc, "一、虫情测报灯采集图像", level=2)
    if insect_images:
        # Group by device_code
        by_device: dict[str, list[dict]] = defaultdict(list)
        for img in insect_images:
            by_device[img["device_code"]].append(img)

        for device_code, imgs in by_device.items():
            _set_heading(doc, f"设备编号：{device_code}", level=3)
            for idx, img in enumerate(imgs, 1):
                caption = f"图 虫情捕获实景 {idx}   采集时间：{img['time']}"
                _insert_url_image(doc, img["url"], caption, width_inches=5.0)
    else:
        _insert_missing_image_placeholder(
            doc,
            "本监测周期未获取到虫情测报灯采集图像，当前保留附录位置占位。",
        )

    # ---- 孢子捕捉仪 ----
    _set_heading(doc, "二、孢子捕捉仪采集图像", level=2)
    if spore_images:
        by_device_s: dict[str, list[dict]] = defaultdict(list)
        for img in spore_images:
            by_device_s[img["device_code"]].append(img)

        for device_code, imgs in by_device_s.items():
            _set_heading(doc, f"设备编号：{device_code}", level=3)
            for idx, img in enumerate(imgs, 1):
                caption = f"图 孢子采集实景 {idx}   采集时间：{img['time']}"
                _insert_url_image(doc, img["url"], caption, width_inches=5.0)
    else:
        _insert_missing_image_placeholder(
            doc,
            "本监测周期未获取到孢子捕捉仪采集图像，当前保留附录位置占位。",
        )


# ---------------------------------------------------------------------------
# Main entry point
# ---------------------------------------------------------------------------

def generate_docx_report(
    summary: dict,
    ai_analysis: str,
    charts: dict,
    ai_images: dict,
    filepath: str,
    figure_manifest: list[dict] | None = None,
):
    """Generate a Word document with images inserted inline at （见图N） positions,
    followed by a device capture image appendix."""
    doc = Document()
    _enable_update_fields_on_open(doc)

    # Global styles
    for style_name in ['Normal', 'Heading 1', 'Heading 2', 'Heading 3', 'Title']:
        if style_name in doc.styles:
            s = doc.styles[style_name]
            s.font.name = 'Microsoft YaHei'
            s.paragraph_format.space_before = Pt(0)
            s.paragraph_format.space_after = Pt(0)
            s.paragraph_format.line_spacing = 1.0
            if s.element.rPr is not None:
                s.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # Footer with page numbers (Page X of Y)
    section = doc.sections[0]
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _tighten_paragraph(fp)
    run = fp.add_run("第 ")
    _set_cjk_font(run)
    _append_field(fp, "PAGE", "1")
    run = fp.add_run(" / ")
    _set_cjk_font(run)
    _append_field(fp, "NUMPAGES", "1")
    run = fp.add_run(" 页")
    _set_cjk_font(run)

    # Title
    period = summary.get("period", {})
    t_start = period.get("start", "")
    t_end = period.get("end", "")
    _add_cover_title(doc, "橡胶林近自然化改造生态效益评估报告")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _tighten_paragraph(p)
    r = p.add_run("三亚市天涯区橡胶林近自然化改造和农田提升监测平台")
    _set_cjk_font(r, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _tighten_paragraph(p)
    r = p.add_run(f"统计周期：{t_start} 至 {t_end}")
    _set_cjk_font(r)

    # Cover image
    cover_b64 = ai_images.get("cover")
    if cover_b64:
        try:
            raw = cover_b64.split(',')[-1]
            img_bytes = base64.b64decode(raw)
            cover_para = doc.add_paragraph()
            cover_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _tighten_paragraph(cover_para)
            cover_para.add_run().add_picture(io.BytesIO(img_bytes), width=Inches(6.0))
        except Exception:
            pass

    doc.add_page_break()

    # Build figure map
    figures = _build_figure_map(summary, charts, ai_images, figure_manifest)

    FIG_REF_RE = re.compile(r'[（(](?:见)?图(\d+)[）)]')
    SUBSECTION_REF_RE = re.compile(r'^\*\*(\d+\.\d+\s+.+?)\*\*$')
    next_expected_fig = 1

    def _insert_figures_through(target_fig: int) -> None:
        nonlocal next_expected_fig
        while next_expected_fig <= target_fig:
            if next_expected_fig in figures:
                b64, caption = figures[next_expected_fig]
                _insert_figure(doc, b64, caption)
                figures.pop(next_expected_fig, None)
            next_expected_fig += 1

    def _strip_inline_markdown(text: str) -> str:
        return re.sub(r"\*\*(.+?)\*\*", r"\1", text or "").strip()

    # Render AI analysis
    if ai_analysis:
        for line in ai_analysis.split("\n"):
            line = line.strip()
            if not line:
                continue

            subsection_match = SUBSECTION_REF_RE.fullmatch(line)
            if line.startswith("## "):
                _set_heading(doc, _strip_inline_markdown(line[3:]), level=1)
            elif line.startswith("### "):
                _set_heading(doc, _strip_inline_markdown(line[4:]), level=2)
            elif subsection_match:
                _set_heading(doc, subsection_match.group(1).strip(), level=2)
            elif line.startswith(("- ", "* ", "• ")):
                clean = _strip_inline_markdown(line[2:])
                _add_paragraph(doc, f"■ {clean}", first_indent=False, left_indent=True)
            else:
                _add_paragraph(doc, _strip_inline_markdown(line))

            # Check for figure references and insert figures immediately after the paragraph
            refs = sorted([int(match.group(1)) for match in FIG_REF_RE.finditer(line)])
            if refs:
                # Insert all figures up to the highest one mentioned in this line
                _insert_figures_through(max(refs))

    # Any remaining AI/chart figures → standard appendix
    if figures:
        _set_heading(doc, "附录：数据图表", level=1)
        _insert_figures_through(max(figures))

    # Device capture image appendix (downloads from CDN URLs)
    _append_device_image_section(doc, summary)

    _remove_empty_body_paragraphs(doc)
    os.makedirs(os.path.dirname(os.path.abspath(filepath)), exist_ok=True)
    doc.save(filepath)
    _refresh_word_fields(filepath)
